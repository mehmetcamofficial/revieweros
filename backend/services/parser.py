from __future__ import annotations

import os
from pathlib import Path
from typing import Optional


MAX_CHARS = int(os.getenv("PARSER_MAX_CHARS", "60000"))

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".pdf",
    ".docx",
}


class ParserError(Exception):
    pass


def clean_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines = []
    previous_blank = False

    for raw_line in text.split("\n"):
        line = " ".join(raw_line.split())

        if not line:
            if not previous_blank:
                lines.append("")
            previous_blank = True
            continue

        lines.append(line)
        previous_blank = False

    cleaned = "\n".join(lines).strip()
    return cleaned


def truncate_text(text: str, max_chars: int = MAX_CHARS) -> str:
    if len(text) <= max_chars:
        return text

    truncated = text[:max_chars].rsplit("\n", 1)[0].strip()

    if not truncated:
        truncated = text[:max_chars].strip()

    return (
        truncated
        + "\n\n[ReviewerOS note: The uploaded document was longer than the processing limit. "
        + f"Only the first {max_chars:,} characters were analyzed.]"
    )


def validate_file_path(file_path: str) -> Path:
    path = Path(file_path)

    if not path.exists():
        raise ParserError(f"File does not exist: {file_path}")

    if not path.is_file():
        raise ParserError(f"Path is not a file: {file_path}")

    if path.stat().st_size == 0:
        raise ParserError("Uploaded file is empty.")

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ParserError(
            f"Unsupported file type '{extension}'. Supported types: {supported}"
        )

    return path


def read_text_file(path: Path) -> str:
    encodings = ["utf-8", "utf-8-sig", "latin-1"]

    last_error: Optional[Exception] = None

    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as error:
            last_error = error

    raise ParserError(f"Failed to decode text file: {last_error}")


def extract_pdf_text(path: Path) -> str:
    try:
        import fitz
    except ImportError as error:
        raise ParserError(
            "PDF parsing requires PyMuPDF. Install it with: pip install pymupdf"
        ) from error

    try:
        document = fitz.open(str(path))
    except Exception as error:
        raise ParserError(f"Failed to open PDF file: {error}") from error

    if document.page_count == 0:
        document.close()
        raise ParserError("PDF has no pages.")

    pages = []

    try:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)

            page_text = page.get_text("text") or ""
            page_text = clean_text(page_text)

            if page_text:
                pages.append(f"\n\n--- Page {page_index + 1} ---\n{page_text}")

    except Exception as error:
        raise ParserError(f"Failed while extracting text from PDF: {error}") from error

    finally:
        document.close()

    text = "\n".join(pages).strip()

    if not text:
        raise ParserError(
            "No readable text could be extracted from this PDF. "
            "It may be scanned/image-based. OCR is not enabled yet."
        )

    return text


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as error:
        raise ParserError(
            "DOCX parsing requires python-docx. Install it with: pip install python-docx"
        ) from error

    try:
        document = Document(str(path))
    except Exception as error:
        raise ParserError(f"Failed to open DOCX file: {error}") from error

    parts = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text or "")

        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        rows = []

        for row in table.rows:
            cells = [clean_text(cell.text or "") for cell in row.cells]
            cells = [cell for cell in cells if cell]

            if cells:
                rows.append(" | ".join(cells))

        if rows:
            parts.append(f"\n--- Table {table_index} ---")
            parts.extend(rows)

    text = "\n".join(parts).strip()

    if not text:
        raise ParserError("No readable text could be extracted from this DOCX file.")

    return text


def extract_text_from_file(file_path: str) -> str:
    path = validate_file_path(file_path)
    extension = path.suffix.lower()

    if extension in {".txt", ".md"}:
        raw_text = read_text_file(path)

    elif extension == ".pdf":
        raw_text = extract_pdf_text(path)

    elif extension == ".docx":
        raw_text = extract_docx_text(path)

    else:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ParserError(
            f"Unsupported file type '{extension}'. Supported types: {supported}"
        )

    cleaned = clean_text(raw_text)

    if not cleaned:
        raise ParserError("No readable text could be extracted from the uploaded file.")

    return truncate_text(cleaned)